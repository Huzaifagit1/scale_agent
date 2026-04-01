from __future__ import annotations

from pathlib import Path
import re

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).parent
SOURCE = ROOT / "API_SOLUTION_BOOK.md"
OUTPUT = ROOT / "API_SOLUTION_BOOK.docx"


def set_cell_shading(paragraph, fill: str) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    p_pr.append(shd)


def set_paragraph_border(paragraph, color: str = "D9DDE7") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        elem = OxmlElement(f"w:{edge}")
        elem.set(qn("w:val"), "single")
        elem.set(qn("w:sz"), "6")
        elem.set(qn("w:space"), "4")
        elem.set(qn("w:color"), color)
        p_bdr.append(elem)
    p_pr.append(p_bdr)


def ensure_style(document: Document, name: str, style_type=WD_STYLE_TYPE.PARAGRAPH):
    try:
        return document.styles[name]
    except KeyError:
        return document.styles.add_style(name, style_type)


def add_field_run(paragraph, field_code: str):
    run = paragraph.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_begin)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)
    return run


def configure_styles(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.font.size = Pt(10)
    normal.font.color.rgb = RGBColor(0x1C, 0x28, 0x33)
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0

    title = document.styles["Title"]
    title.font.name = "Arial"
    title._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = RGBColor(0x1C, 0x28, 0x33)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Arial"
    heading_1._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_1.font.size = Pt(14)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)
    heading_1.paragraph_format.space_before = Pt(18)
    heading_1.paragraph_format.space_after = Pt(6)

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Arial"
    heading_2._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_2.font.size = Pt(12)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)
    heading_2.paragraph_format.space_before = Pt(12)
    heading_2.paragraph_format.space_after = Pt(4)

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Arial"
    heading_3._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    heading_3.font.size = Pt(11)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(0x1C, 0x28, 0x33)
    heading_3.paragraph_format.space_before = Pt(8)
    heading_3.paragraph_format.space_after = Pt(3)

    bullet = ensure_style(document, "Book Bullet")
    bullet.base_style = normal
    bullet.font.name = "Arial"
    bullet._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    bullet.font.size = Pt(10)
    bullet.font.color.rgb = RGBColor(0x1C, 0x28, 0x33)
    bullet.paragraph_format.left_indent = Inches(0.25)
    bullet.paragraph_format.space_after = Pt(2)

    code = ensure_style(document, "Code Block")
    code.base_style = normal
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
    code.font.size = Pt(9)
    code.paragraph_format.left_indent = Inches(0.22)
    code.paragraph_format.right_indent = Inches(0.12)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(3)


INLINE_CODE_RE = re.compile(r"`([^`]+)`")


def add_inline_runs(paragraph, text: str) -> None:
    last = 0
    for match in INLINE_CODE_RE.finditer(text):
        if match.start() > last:
            paragraph.add_run(text[last:match.start()])
        run = paragraph.add_run(match.group(1))
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x8B, 0x1E, 0x3F)
        last = match.end()
    if last < len(text):
        paragraph.add_run(text[last:])


def render_markdown(document: Document, markdown: str) -> None:
    lines = markdown.splitlines()
    in_code = False
    code_buffer: list[str] = []
    title_done = False

    for raw_line in lines:
        line = raw_line.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_buffer = []
            else:
                paragraph = document.add_paragraph(style="Code Block")
                code_text = "\n".join(code_buffer)
                paragraph.add_run(code_text)
                set_cell_shading(paragraph, "F6F8FB")
                set_paragraph_border(paragraph)
                in_code = False
                code_buffer = []
            continue

        if in_code:
            code_buffer.append(raw_line)
            continue

        stripped = line.strip()
        if not stripped:
            document.add_paragraph("")
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not title_done:
                paragraph = document.add_paragraph(style="Title")
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                paragraph.add_run(text)
                title_done = True
            else:
                paragraph = document.add_paragraph(style="Heading 1")
                paragraph.add_run(text)
            continue

        if stripped.startswith("## "):
            paragraph = document.add_paragraph(style="Heading 1")
            paragraph.add_run(stripped[3:].strip())
            continue

        if stripped.startswith("### "):
            paragraph = document.add_paragraph(style="Heading 2")
            paragraph.add_run(stripped[4:].strip())
            continue

        if stripped.startswith("#### "):
            paragraph = document.add_paragraph(style="Heading 3")
            paragraph.add_run(stripped[5:].strip())
            continue

        if stripped.startswith("- "):
            paragraph = document.add_paragraph(style="Book Bullet")
            paragraph.style = document.styles["List Bullet"]
            add_inline_runs(paragraph, stripped[2:].strip())
            continue

        paragraph = document.add_paragraph(style="Normal")
        paragraph.paragraph_format.space_after = Pt(0)
        add_inline_runs(paragraph, stripped)


def add_reference_like_header_footer(document: Document) -> None:
    section = document.sections[0]

    header = section.header.paragraphs[0]
    header.text = ""
    left = header.add_run("API Solution Book  |  Project Requirements Document")
    left.font.name = "Arial"
    left._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    left.font.size = Pt(9)
    left.font.color.rgb = RGBColor(0xAA, 0xB7, 0xB8)

    tab = header.add_run("\t")
    tab.font.size = Pt(9)

    right = header.add_run("CONFIDENTIAL")
    right.font.name = "Arial"
    right._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    right.font.size = Pt(9)
    right.font.bold = True
    right.font.color.rgb = RGBColor(0x2E, 0x86, 0xC1)

    footer = section.footer.paragraphs[0]
    footer.text = ""
    lead = footer.add_run("Prepared for Internal Use Only  |  ")
    lead.font.name = "Arial"
    lead._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    lead.font.size = Pt(8)
    lead.font.color.rgb = RGBColor(0xAA, 0xB7, 0xB8)

    page_label = footer.add_run("Page ")
    page_label.font.name = "Arial"
    page_label._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    page_label.font.size = Pt(8)
    page_label.font.color.rgb = RGBColor(0xAA, 0xB7, 0xB8)

    page_run = add_field_run(footer, " PAGE ")
    page_run.font.name = "Arial"
    page_run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    page_run.font.size = Pt(8)
    page_run.font.color.rgb = RGBColor(0xAA, 0xB7, 0xB8)


def main() -> None:
    markdown = SOURCE.read_text(encoding="utf-8")
    document = Document()
    configure_styles(document)
    add_reference_like_header_footer(document)
    render_markdown(document, markdown)
    document.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
